import os

import docx


def fix_answer_numeric(source_dir, dest_dir):
    for root, dirs, files in os.walk(source_dir):
        for file in files:
            doc = docx.Document(os.path.join(root, file))
            last_table = doc.tables[len(doc.tables) - 1]
            for i in range(1, 26):
                try:
                    print()
                    last_table.rows[i].cells[1].text = chr(ord("А") - 1 + int(last_table.rows[i].cells[1].text))
                    last_table.rows[i].cells[1].paragraphs[0].style = "!СКИФ-Ключ-Ответ"
                except:
                    print(last_table.rows[i].cells[1].text)
            dest_root = root.replace(source_dir, dest_dir)
            if not os.path.exists(dest_root):
                os.makedirs(dest_root)
            doc.save(os.path.join(dest_root, file))


# fix_answer_numeric("./stasik", "./stasik2")

def to_gift2(from_file, to_file):
    with open(from_file, "r", encoding="utf-8") as f:
        a = f.read()

    a = a.strip()
    while "\n\n\n" in a:
        a = a.replace("\n\n\n", "\n\n")

    q = a.split("\n\n")

    def to_quest(arr):
        q = f"::{arr['num']}.:: {arr['text']}\n{'{'}\n"
        for i in range(len(arr['rows'])):
            q += "    = " + arr['rows'][i] + "\n"
        return q + "}"

    res = []
    start = 1
    for i in range(len(q)):
        t = q[i].split("\n")
        arr = {
            "num": start + i,
            "text": t[0].split(" ", 1)[-1],
            "rows": list(map(lambda x: x.strip(), t[1].split(",")))
        }
        res.append(arr)

    with open(to_file, "w", encoding="utf-8") as f:
        f.write("\n\n".join(map(to_quest, res)))


def to_gift1(from_file, to_file):
    with open(from_file, "r", encoding="utf-8") as f:
        a = f.read()

    a = a.strip()
    while "\n\n\n" in a:
        a = a.replace("\n\n\n", "\n\n")

    q = a.split("\n\n")

    def to_quest(arr):
        q = f"::{arr['num']}.:: {arr['text']}\n{'{'}\n"
        for i in range(len(arr['rows'])):
            q += ("    = " if arr['rows'][i].startswith("*") else "    ~ ") + (
                arr['rows'][i][1:] if arr['rows'][i].startswith("*") else arr['rows'][i]) + "\n"
        return q + "}"

    res = []
    start = 1
    for i in range(len(q)):
        t = q[i].split("\n")
        arr = {
            "num": start + i,
            "text": t[0].split(" ", 1)[-1],
            "rows": []
        }
        for j in range(1, len(t)):
            r = []
            arr["rows"].append(t[j].split(" ", 1)[-1].strip())
        res.append(arr)

    with open(to_file, "w", encoding="utf-8") as f:
        f.write("\n\n".join(map(to_quest, res)))


def to_gift3(from_file, to_file):
    doc = docx.Document(from_file)

    start = 1
    out = []
    for i in range(len(doc.tables)):
        table = doc.tables[i]
        arr = {
            "num": start + i,
            "text": "",
            "rows": []
        }
        for j in range(len(table.rows)):
            r = []
            for k in range(len(table.rows[j].cells)):
                r.append(table.rows[j].cells[k].text)
            arr["rows"].append(" -> ".join(r))
        out.append(f"::{arr['num']}.:: {arr['text']}\n{'{'}\n    = " + "\n    = ".join(arr['rows']) + "\n}")
    with open(to_file, "w", encoding="utf-8") as f:
        f.write("\n\n".join(out))


# Closed
#to_gift1("Z:/Desktop/Мобильные приложения4.txt", "Z:/Desktop/Васильев - Мобильные приложения - закрытые.gift")

# opened
#to_gift2("Z:/Desktop/Мобильные приложения3.txt", "Z:/Desktop/Васильев - Мобильные приложения - открытые.gift")

# to_gift2("./source/any/from33.txt", "./to33.txt")
# Последовательность
to_gift3("Z:/Desktop/Таблица 13.docx", "Z:/Desktop/Васильев - Мобильные приложения - соответствие.gift")
