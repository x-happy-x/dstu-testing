import json
import os

import docx
import docxtpl

from app.utils import discipline_filter, get_temp_filename
from app.parser.plan import load_plan


def create_docx(plan_info, new_plan_path, template,
                shorten_disciplines=[], blacklist_disciplines=[], other_column={}, table_text_style='TableText'):
    temp_file = get_temp_filename('docx')

    doc = docx.Document(template)
    tables = {
        'УК': doc.tables[-3],
        'ОПК': doc.tables[-2],
        'ПК': doc.tables[-1],
    }
    start_columns = {
        'УК': 1,
        'ОПК': 0,
        'ПК': 1,
    }

    comp_dict = plan_info['competencies']
    for comp in comp_dict:
        table = tables[comp]
        row_idx = 1
        col_idx = start_columns[comp]
        competencies = comp_dict[comp]
        for competence in competencies:
            # print(competence)
            indicators = competencies[competence]['sub']
            first_cell1 = -1
            last_cell1 = -1
            for indicator in indicators:
                # print(indicator)
                disciplines = indicators[indicator]['sub']
                table.add_row()
                row_idx += 1
                if first_cell1 == -1:
                    first_cell1 = row_idx
                    table.cell(first_cell1, col_idx).paragraphs[
                        0].text = f"{competence} {competencies[competence]['text']}"

                last_cell1 = row_idx
                table.cell(last_cell1, col_idx + 1).paragraphs[0].text = f"{indicator} {indicators[indicator]['text']}"
                table.cell(last_cell1, col_idx + 2).paragraphs[0].text = "\n".join(
                    discipline_filter(disciplines, shorten_disciplines, blacklist_disciplines))

                for i in range(len(table.rows[row_idx].cells)):
                    table.cell(row_idx, i).paragraphs[0].style = table_text_style

            table.cell(first_cell1, col_idx).merge(table.cell(last_cell1, col_idx))
            if comp in other_column and competence.split('-')[-1] in other_column[comp]:
                table.cell(first_cell1, col_idx - 1).paragraphs[
                    0].text = f"{other_column[comp][competence.split('-')[-1]]}"
                table.cell(first_cell1, col_idx - 1).merge(table.cell(last_cell1, col_idx - 1))
    doc.save(temp_file)
    doc = docxtpl.DocxTemplate(temp_file)
    doc.render(plan_info['info'])
    doc.save(new_plan_path)
    os.remove(temp_file)


def excel_to_docx(source, dest, template, json_save=False):
    groups = os.listdir(source)
    txt = {}
    for group in groups:
        group_path = os.path.join(source, group)
        plans = list(filter(lambda x: not x.startswith("~") and x.endswith(".xlsx"), os.listdir(group_path)))

        new_group_path = os.path.join(dest, group)
        if os.path.exists(new_group_path) and not os.path.isdir(new_group_path):
            os.remove(new_group_path)
        if not os.path.exists(new_group_path):
            os.makedirs(new_group_path)

        for plan in plans:
            plan_path = os.path.join(group_path, plan)
            plan_info = load_plan(group, plan_path)
            if json_save:
                if group not in txt:
                    txt[group] = {}
                txt[group][plan_info['info']['year']] = plan_info
            new_plan_path = os.path.join(new_group_path, plan_info['filename'])
            create_docx(plan_info, new_plan_path, template)

    if json_save:
        json_path = os.path.join(dest, json_save)

        json_dir = os.path.dirname(json_path)
        if not os.path.exists(json_dir):
            os.makedirs(json_dir)

        with open(json_path, 'w', encoding='utf-8') as f:
            f.write(json.dumps(txt, ensure_ascii=False).encode('utf-8').decode())

    return txt


def json_to_docx(source="", dest="", template="", plan_data: dict[str] = None):
    if plan_data is None:
        with open(source, 'r', encoding='utf-8') as f:
            plan_data = json.load(f)

    for group in plan_data:
        new_group_path = os.path.join(dest, group)
        if os.path.exists(new_group_path) and not os.path.isdir(new_group_path):
            os.remove(new_group_path)
        if not os.path.exists(new_group_path):
            os.makedirs(new_group_path)
        for plan in plan_data[group]:
            plan_info = plan_data[group][plan]
            new_plan_path = os.path.join(new_group_path, plan_info['filename'])
            # print(group, plan)
            create_docx(template, plan_info, new_plan_path)

    return plan_data


def json_split(source="", dest="", stat=True, diff=True, plan_data: dict[str] = None):
    if plan_data is None:
        with open(source, 'r', encoding='utf-8') as f:
            plan_data = json.load(f)

    for group in plan_data:
        new_group_path = os.path.join(dest, group)
        if os.path.exists(new_group_path) and not os.path.isdir(new_group_path):
            os.remove(new_group_path)
        if not os.path.exists(new_group_path):
            os.makedirs(new_group_path)
        for plan in plan_data[group]:
            plan_info = plan_data[group][plan]
            new_plan_path = os.path.join(new_group_path, str(plan_info['info']['year']) + ".json")
            # print(group, plan)

            if stat:
                list_comp = {}
                list_indicators = {}
                list_disciplines = {}
                for comp in plan_info['competencies']:
                    competencies = plan_info['competencies'][comp]
                    for competence in competencies:
                        competence_ = f"{competence} {competencies[competence]['text']}"
                        if competence_ not in list_comp:
                            list_comp[competence_] = 1
                        else:
                            list_comp[competence_] += 1
                        indicators = competencies[competence]['sub']
                        for indicator in indicators:
                            indicator_ = f"{indicator} {indicators[indicator]['text']}"
                            if indicator_ not in list_indicators:
                                list_indicators[indicator_] = 1
                            else:
                                list_indicators[indicator_] += 1
                            for discipline in discipline_filter(
                                    disciplines=indicators[indicator]['sub'],
                                    shorten=['Основы проектной деятельности', 'Иностранный язык'],
                                    blacklist=[]
                            ):
                                if discipline not in list_disciplines:
                                    list_disciplines[discipline] = 1
                                else:
                                    list_disciplines[discipline] += 1
                plan_info['stat'] = {
                    'Компетенции': sorted(list_comp),
                    "Индикаторы": sorted(list_indicators),
                    "Дисциплины": {k: v for k, v in sorted(list_disciplines.items())}
                }

            with open(new_plan_path, 'w', encoding='utf-8') as f:
                f.write(json.dumps(plan_info, ensure_ascii=False).encode('utf-8').decode())

    return plan_data
