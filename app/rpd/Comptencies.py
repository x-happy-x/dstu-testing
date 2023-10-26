from typing import List
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Cm
from docx.shared import Pt

from docx.table import _Cell
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def set_margins(section, top=None, bottom=None, left=None, right=None, all=None):
    if all is not None:
        if top is None: section.top_margin = Cm(all)
        if bottom is None: section.bottom_margin = Cm(all)
        if left is None: section.left_margin = Cm(all)
        if right is None: section.right_margin = Cm(all)
    if top is not None:
        section.top_margin = Cm(top)
    if bottom is not None:
        section.bottom_margin = Cm(bottom)
    if left is not None:
        section.left_margin = Cm(left)
    if right is not None:
        section.right_margin = Cm(right)


def rotate(section):
    h, w = section.page_height, section.page_width
    orient = section.orientation
    if orient == WD_ORIENT.PORTRAIT:
        section.orientation = WD_ORIENT.LANDSCAPE
    else:
        section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = h
    section.page_height = w


def set_cell_border(cell: _Cell, **kwargs):
    """
    Set cell`s border
    Usage:
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))


def cell_center(cell):
    for p in cell.paragraphs:
        p.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def cell_bold(cell):
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True


def cell_border(cell):
    border = {"sz": 6, "color": "#000000", "val": "single"}
    set_cell_border(
        cell,
        top=border,
        bottom=border,
        start=border,
        end=border,
    )


def cell_body(doc, cell):
    for p in cell.paragraphs:
        p.style = doc.styles['TableBodyStyle']
        p.paragraph_format.space_after = Pt(0)


def cell_header(doc, cell):
    for p in cell.paragraphs:
        p.style = doc.styles['TableHeaderStyle']
        p.paragraph_format.space_after = Pt(0)


def add_table_header_text(doc, text, font_size=14, alignment=WD_TABLE_ALIGNMENT.CENTER, style="Normal"):
    p_header = doc.add_paragraph(text, style=style)
    p_header.alignment = alignment
    p_header.paragraph_format.space_after = Pt(0)
    p_header.runs[0].font.size = Pt(font_size)


def add_table_comps(doc, cb, headers, level_names, cells_text, merge_cols, cols_width, styles, table=None):
    if table is None:
        table = doc.add_table(rows=1, cols=7)

    row = table.rows[0]
    row.cells[0].text = headers['comp']
    row.cells[1].text = headers['indi_type']
    row.cells[2].text = headers['indi']
    row.cells[3].text = headers['type']
    row.cells[4].text = headers['control']
    row.cells[5].text = headers['OM']
    row.cells[6].text = headers['crit']

    for style in styles:
        if 'header' not in style['type']:
            continue
        for ind in style['indexes']:
            style['func'](row.cells[ind])

    table.add_row()

    row_index = 1
    for competence in cb.competencies():
        for indicator in competence.indicators():

            last_index = row_index
            row = table.rows[row_index]

            row.cells[0].text = indicator.indi_code
            row.cells[3].text = cells_text['type']
            row.cells[4].text = cells_text['control']
            row.cells[5].text = cells_text['OM']
            row.cells[6].text = cells_text['crit']

            for level in indicator.levels():
                if level.contents is None or len(level.contents) < 4:
                    continue
                row.cells[1].text = level_names[level.level_id % len(level_names)]
                row.cells[2].text = level.contents

                row_index += 1
                row = table.add_row()

            row1 = table.rows[last_index]
            row2 = table.rows[row_index - 1]

            for row_id in range(last_index, row_index):
                for style in styles:
                    if 'body' not in style['type']:
                        continue
                    for ind in style['indexes']:
                        style['func'](table.rows[row_id].cells[ind])

            for m_id in merge_cols:
                row1.cells[m_id].merge(row2.cells[m_id])

    row._element.getparent().remove(row._element)

    for i in range(len(table.columns)):
        for cell in table.columns[i].cells:
            cell.width = Cm(cols_width[i])

    return table


class Competence:
    card_id: int
    comp_id: int
    comp_code: str
    comp_description: str
    indicator_list: List

    def __init__(self, cb, **kwargs):
        self.cb = cb
        self.comp_id = kwargs.get('compId')
        self.card_id = kwargs.get('cardId')
        self.comp_code = kwargs.get('compCode')
        self.comp_description = kwargs.get('compDescription')

        self.indicator_list = []

    def set(self, comp_id: int, card_id: int, comp_code: str, comp_description: str) -> None:
        self.comp_id = comp_id
        self.card_id = card_id
        self.comp_code = comp_code
        self.comp_description = comp_description

    def add_indicator(self, indicator=None, **kwargs):
        if indicator is not None:
            self.indicator_list.append(indicator)
        else:
            self.indicator_list.append(Indicator(self, **kwargs))

    def indicators(self):
        return self.indicator_list


class Indicator:
    competence: Competence
    indicator_id: int
    indi_code: str
    indi_description: str
    level_list: List

    def __init__(self, competence: Competence, **kwargs):
        self.competence = competence
        self.indicator_id = kwargs.get('indicatorId')
        self.indi_code = kwargs.get('indiCode')
        self.indi_description = kwargs.get('indiDescription')
        self.level_list = []

    def set(self, indicator_id: int, indi_code: str, indi_description: str):
        self.indicator_id = indicator_id
        self.indi_code = indi_code
        self.indi_description = indi_description

    def add_level(self, level=None, **kwargs):
        if level is not None:
            self.level_list.append(level)
        else:
            self.level_list.append(Level(self, **kwargs))

    def levels(self):
        return self.level_list


class Level:
    indicator: Indicator
    category_id: int
    category_name: str
    level_id: int
    is_mock: bool
    contents: str
    rating_scale: str
    mark_scale: str
    code: str
    mark5: None
    mark4: None
    mark3: None
    mark2: None
    task1: None
    task2: None
    fos_inter_control_task1: None
    fos_inter_control_task2: None
    fos_current_control_task1: None
    fos_current_control_task2: None
    download_fos_url: str
    upload_fos_url: str
    delete_fos_url: str
    has_current_fos: None
    has_intermediate_fos: None

    def __init__(self, indicator: Indicator, **kwargs):
        self.indicator = indicator
        self.category_id = kwargs.get('category_id')
        self.category_name = kwargs.get('categoryName')
        self.level_id = kwargs.get('levelId')
        self.is_mock = kwargs.get('isMock')
        self.contents = kwargs.get('contents')
        self.rating_scale = kwargs.get('ratingScale')
        self.mark_scale = kwargs.get('markScale')
        self.code = kwargs.get('code')
        self.mark5 = kwargs.get('mark5')
        self.mark4 = kwargs.get('mark4')
        self.mark3 = kwargs.get('mark3')
        self.mark2 = kwargs.get('mark2')
        self.task1 = kwargs.get('task1')
        self.task2 = kwargs.get('task2')
        self.fos_inter_control_task1 = kwargs.get('fos_interControl_task1')
        self.fos_inter_control_task2 = kwargs.get('fos_interControl_task2')
        self.fos_current_control_task1 = kwargs.get('fos_currentControl_task1')
        self.fos_current_control_task2 = kwargs.get('fos_currentControl_task2')
        self.download_fos_url = kwargs.get('downloadFosUrl')
        self.upload_fos_url = kwargs.get('uploadFosUrl')
        self.delete_fos_url = kwargs.get('deleteFosUrl')
        self.has_current_fos = kwargs.get('hasCurrentFos')
        self.has_intermediate_fos = kwargs.get('hasIntermediateFos')

    def set(self, category_id: int, category_name: str, level_id: int, is_mock: bool, contents: str, rating_scale: str,
            mark_scale: str, code: str, mark5: None, mark4: None, mark3: None, mark2: None, task1: None, task2: None,
            fos_inter_control_task1: None, fos_inter_control_task2: None, fos_current_control_task1: None,
            fos_current_control_task2: None, download_fos_url: str, upload_fos_url: str, delete_fos_url: str,
            has_current_fos: None, has_intermediate_fos: None) -> None:
        self.category_id = category_id
        self.category_name = category_name
        self.level_id = level_id
        self.is_mock = is_mock
        self.contents = contents
        self.rating_scale = rating_scale
        self.mark_scale = mark_scale
        self.code = code
        self.mark5 = mark5
        self.mark4 = mark4
        self.mark3 = mark3
        self.mark2 = mark2
        self.task1 = task1
        self.task2 = task2
        self.fos_inter_control_task1 = fos_inter_control_task1
        self.fos_inter_control_task2 = fos_inter_control_task2
        self.fos_current_control_task1 = fos_current_control_task1
        self.fos_current_control_task2 = fos_current_control_task2
        self.download_fos_url = download_fos_url
        self.upload_fos_url = upload_fos_url
        self.delete_fos_url = delete_fos_url
        self.has_current_fos = has_current_fos
        self.has_intermediate_fos = has_intermediate_fos


class CompetenceBoard:
    competence_list: List[Competence]

    def __init__(self, rp):
        self.rp = rp
        self.competence_list = []

    def generate_table1(self, save=None, doc=None, table=None):
        if doc is None:
            doc = Document()

        current_section = doc.sections[-1]
        rotate(current_section)
        set_margins(current_section, all=1, top=2)

        headers = {
            'comp': 'Код компетенции',
            'indi_type': 'Уровень освоения компетенции',
            'indi': 'Дескрипторы компетенции (результаты обучения, показатели достижения результата обучения, '
                    'которые обучающийся может продемонстрировать)',
            'type': 'Виды работ в рамках практики, формы и методы обучения, способствующие формированию и развитию '
                    'компетенции',
            'control': 'Контролируемые разделы и темы практики',
            'OM': 'Оценочные материалы (оценочные средства), используемые для оценки уровня сформированности '
                  'компетенции',
            'crit': 'Способы оценивания компетенций',
        }

        level_names = [
            'Запоминание',
            'Понимание',
            'Применение',
        ]

        cells_text = {
            'type': 'Контактная работа с преподавателем, самостоятельная работа',
            'control': 'Разделы 1-3',
            'OM': 'Отчет по практике, индивидуальное задание на практику, защита отчета по практике',
            'crit': 'выполнение индивидуального задания на практику, выполнение отчета в назначенный срок, ответы на '
                    'вопросы преподавателя по отчету по практике'
        }

        merge_cols = [0, 3, 4, 5, 6]
        cols_width = [1, 1, 25, 1, 1, 1, 1]
        styles = [
            {
                'indexes': [0, 1, 3, 4, 5, 6],
                'type': 'header,body',
                'func': cell_center
            },
            {
                'indexes': [2],
                'type': 'header',
                'func': cell_center
            },
            {
                'indexes': [1],
                'type': 'body',
                'func': cell_bold
            },
            {
                'indexes': range(7),
                'type': 'header,body',
                'func': cell_border
            },
            {
                'indexes': range(7),
                'type': 'header',
                'func': lambda x: cell_header(doc, x)
            },
            {
                'indexes': range(7),
                'type': 'body',
                'func': lambda x: cell_body(doc, x)
            }
        ]
        p_header_text = "Таблица 1 - Формирование компетенций в процессе прохождения практики"

        style = doc.styles.add_style('TableBodyStyle', WD_STYLE_TYPE.PARAGRAPH)
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(11)

        style = doc.styles.add_style('TableHeaderStyle', WD_STYLE_TYPE.PARAGRAPH)
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(11)

        add_table_header_text(doc, p_header_text)
        table = add_table_comps(doc, self, headers, level_names, cells_text, merge_cols, cols_width, styles, table)
        if save is not None:
            doc.save(save)
        return table

    def add_item(self, item):
        comp: Competence
        indicator: Indicator
        cmp: Competence | None = None

        # Поиск компетенции в имеющихся
        for comp in self.competence_list:
            if comp.comp_id == item['compId']:
                cmp = comp
                break

        # Если нет добавление нового
        if cmp is None:
            cmp = Competence(self, **item)
            self.competence_list.append(cmp)

        ind: Indicator | None = None
        # Поиск компетенции в имеющихся
        for indicator in cmp.indicators():
            if indicator.indicator_id == item['indicatorId']:
                ind = indicator
                break

        # Если нет добавление нового
        if ind is None:
            cmp.add_indicator(**item)
            ind = cmp.indicators()[-1]

        # Если уровень пустой не добавляем ничего
        if item['contents'] is None or len(item['contents'].strip()) < 3:
            return

        # Иначе добавляем новый уровень
        ind.add_level(**item)

    def competencies(self):
        return self.competence_list
