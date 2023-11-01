import os
import re
from datetime import datetime

import docxtpl
from docx import Document

from app.entity import *
from app.parser import gift, Gift
from app.parser.doc.editor import (
    find_paragraph_id,
    insert_paragraph_after,
    insert_html_after,
    insert_markdown_after,
    add_multiple_question,
    add_matching_question,
    add_short_question,
    contains_html,
    copy_table_after,
    text2latex,
    get_latex_equations,
    move_table_after,
    remove_paragraph,
    clean_string_regex,
    iter_block_items,
    insert_text_in,
    insert_text_after,
    replace_text,

    __PATTERN_FOR_NUMERATION_REGEX__,
)


class Layout:
    struct: dict
    groups: dict[str, dict]
    stats: dict[str, dict]

    def __init__(self, template, struct=None, path=None, styles=None, html_convert=True):
        if styles is None:
            styles = STYLES
        if struct is not None:
            self.set_struct(struct)

        self.document = Document(template)
        self.template = template

        self.path = path
        self.styles = styles

        self.skipped = []
        self.added = []
        self.changed = False

        self.html_convert = html_convert

        self.add_quest_func = {
            MULTIPLE_CHOICE_RADIO: self.add_multi_radio,
            MULTIPLE_CHOICE_CHECKBOX: self.add_multi_choice,
            MATCHING: self.add_matching,
            SHORT: self.add_short,
        }

    def set_struct(self, struct):
        self.struct = struct
        self.groups = {s['group']: {} for s in struct}
        self.stats = {k: {} for k in self.groups}

    def add_prefix(self, paragraph, question, prefix=None):

        self.changed = True
        self.added.append(question)

        if prefix is not None and len(prefix.strip()) > 0:
            paragraph = insert_markdown_after(paragraph, prefix, self.styles['question'])[-1]
            paragraph = insert_paragraph_after(paragraph, "", self.styles['title'])

        return paragraph

    def add_postfix(self, paragraph, question, postfix=None):

        if postfix is not None and len(postfix.strip()) > 0:
            paragraph = insert_markdown_after(paragraph, postfix, self.styles['question'])[-1]
            paragraph = insert_paragraph_after(paragraph, "", self.styles['title'])

        return paragraph

    def get_paragraph(self, group, question):

        category = question.get_category()
        cat = category.split(" (")[0].strip()

        if category not in self.groups[group]:
            if len(self.groups[group]) == 0:
                paragraph_id = find_paragraph_id(
                    self.document,
                    "{{ " + group + " }}",
                    FindFunc.EQUALS,
                    0
                )
                paragraph = self.document.paragraphs[paragraph_id]
                paragraph.text = category
                paragraph.style = self.styles["title"]
                paragraph = insert_paragraph_after(paragraph, "", self.styles['title'])
                self.groups[group][category] = paragraph
            else:
                paragraph = self.groups[group][list(self.groups[group].keys())[-1]]
                paragraph = insert_paragraph_after(paragraph, category, self.styles['title'])
                paragraph = insert_paragraph_after(paragraph, "", self.styles['title'])
                self.groups[group][category] = paragraph
            self.stats[group][cat] = 1
        else:
            paragraph = self.groups[group][category]
            self.stats[group][cat] += 1

        return [paragraph, category, cat]

    def add_multi_radio(self, group, question, prefix=None, postfix=None):
        if group not in self.groups:
            return

        paragraph, category, cat = self.get_paragraph(group, question)
        paragraph = self.add_prefix(paragraph, question, prefix)

        text = question.get_text()
        if self.html_convert and not text.startswith("[html]") and contains_html(text):
            text = f"[html]{text}"

        paragraph = add_multiple_question(
            paragraph,
            question,
            text,
            STYLES,
        )

        paragraph = self.add_postfix(paragraph, question, postfix)
        self.groups[group][category] = paragraph

        return paragraph

    def add_multi_choice(self, group, question, prefix=None, postfix=None):
        return self.add_multi_radio(group, question, prefix, postfix)

    def add_short(self, group, question, prefix=None, postfix=None):
        if group not in self.groups:
            return

        paragraph, category, cat = self.get_paragraph(group, question)
        paragraph = self.add_prefix(paragraph, question, prefix)

        text = question.get_text()
        if self.html_convert and not text.startswith("[html]") and contains_html(text):
            text = f"[html]{text}"

        paragraph = add_short_question(
            paragraph,
            question,
            text,
            STYLES,
        )

        paragraph = self.add_postfix(paragraph, question, postfix)
        self.groups[group][category] = paragraph

        return paragraph

    def add_matching(self, group, question, prefix=None, postfix=None):
        if group not in self.groups:
            return

        paragraph, category, cat = self.get_paragraph(group, question)
        paragraph = self.add_prefix(paragraph, question, prefix)

        text = question.get_text()
        if self.html_convert and not text.startswith("[html]") and contains_html(text):
            text = f"[html]{text}"

        paragraph = add_matching_question(
            self.document,
            paragraph,
            question,
            text,
            STYLES,
        )

        paragraph = self.add_postfix(paragraph, question, postfix)
        self.groups[group][category] = paragraph

        return paragraph

    def add(self, group, question, prefix=None, postfix=None):
        self.add_quest_func[question.get_type()](group, question, prefix, postfix)

    def find(self, name):
        name = name.lower().strip()
        for question in self.added:
            if question.name.lower().strip() == name:
                return question
        return None

    def build(self, question_id=1, temp_file=None, info=None, to_file=None):
        questions_sorted = []
        for paragraph in self.document.paragraphs:
            match = re.match(__PATTERN_FOR_NUMERATION_REGEX__, paragraph.text)
            if match is not None:
                q_id = match.group(0).replace("num_", "")[2:-2].strip()
                paragraph.text = paragraph.text.replace(match.group(0), str(question_id))
                question_id += 1
                questions_sorted.append(q_id)

        key_table = self.document.tables[-1]

        for row in key_table.rows[1:]:
            row._element.getparent().remove(row._element)

        first_column = len(questions_sorted) // 2

        for i in range(len(questions_sorted)):
            if i < first_column:
                row = key_table.add_row()
                start_column = 0
            else:
                row = key_table.rows[1 + (i % first_column)]
                start_column = 3

            row.cells[start_column].text = f"{i + 1}"
            for p in row.cells[start_column].paragraphs:
                p.style = self.styles["key_num"]

            right = self.find(questions_sorted[i]).get_right_answer()
            right = (right
                     .replace(r'\:', ':')
                     .replace(r'\=', '=')
                     .replace(r'\{', '{')
                     .replace(r'\}', '}')
                     .replace(r'\#', '#')
                     # .lower() # все ответы нижнего регистра
                     .strip())

            row.cells[start_column + 1].text = f"{right}"
            for p in row.cells[start_column + 1].paragraphs:
                p.style = self.styles["key_answer"]

        if temp_file is None:
            temp_file = f'./.cache/LAYOUT_BUILDER/temp-{datetime.now():%Y_%m_%d %H_%M_%S%z}.docx'
            dest = os.path.split(temp_file)[0]
            if len(dest) > 0 and not os.path.exists(dest):
                os.makedirs(dest)

        self.document.save(temp_file)

        context = {
            "simple_c": self.stats['ЗАКРЫТЫЕ_ВОПРОСЫ'][CATEGORY[1]] if CATEGORY[1] in self.stats[
                'ЗАКРЫТЫЕ_ВОПРОСЫ'] else 0,
            "simple_a": self.stats['ВОПРОСЫ_НА_СООТВЕТСТВИЕ'][CATEGORY[1]] if CATEGORY[1] in self.stats[
                'ВОПРОСЫ_НА_СООТВЕТСТВИЕ'] else 0,
            "simple_o": self.stats['ОТКРЫТЫЕ_ВОПРОСЫ'][CATEGORY[1]] if CATEGORY[1] in self.stats[
                'ОТКРЫТЫЕ_ВОПРОСЫ'] else 0,
            "simple_e": self.stats['ЭССЕ'][CATEGORY[1]] if CATEGORY[1] in self.stats['ЭССЕ'] else 0,

            "medium_c": self.stats['ЗАКРЫТЫЕ_ВОПРОСЫ'][CATEGORY[2]] if CATEGORY[2] in self.stats[
                'ЗАКРЫТЫЕ_ВОПРОСЫ'] else 0,
            "medium_a": self.stats['ВОПРОСЫ_НА_СООТВЕТСТВИЕ'][CATEGORY[2]] if CATEGORY[2] in self.stats[
                'ВОПРОСЫ_НА_СООТВЕТСТВИЕ'] else 0,
            "medium_o": self.stats['ОТКРЫТЫЕ_ВОПРОСЫ'][CATEGORY[2]] if CATEGORY[2] in self.stats[
                'ОТКРЫТЫЕ_ВОПРОСЫ'] else 0,
            "medium_e": self.stats['ЭССЕ'][CATEGORY[2]] if CATEGORY[2] in self.stats['ЭССЕ'] else 0,

            "hard_c": self.stats['ЗАКРЫТЫЕ_ВОПРОСЫ'][CATEGORY[3]] if CATEGORY[3] in self.stats[
                'ЗАКРЫТЫЕ_ВОПРОСЫ'] else 0,
            "hard_a": self.stats['ВОПРОСЫ_НА_СООТВЕТСТВИЕ'][CATEGORY[3]] if CATEGORY[2] in self.stats[
                'ВОПРОСЫ_НА_СООТВЕТСТВИЕ'] else 0,
            "hard_o": self.stats['ОТКРЫТЫЕ_ВОПРОСЫ'][CATEGORY[3]] if CATEGORY[3] in self.stats[
                'ОТКРЫТЫЕ_ВОПРОСЫ'] else 0,
            "hard_e": self.stats['ЭССЕ'][CATEGORY[3]] if CATEGORY[3] in self.stats['ЭССЕ'] else 0,
        }

        context["all_c"] = context["simple_c"] + context["medium_c"] + context["hard_c"]
        context["all_a"] = context["simple_a"] + context["medium_a"] + context["hard_a"]
        context["all_o"] = context["simple_o"] + context["medium_o"] + context["hard_o"]
        context["all_e"] = context["simple_e"] + context["medium_e"] + context["hard_e"]

        context["all_simple"] = context["simple_c"] + context["simple_a"] + context["simple_o"] + context["simple_e"]
        context["all_medium"] = context["medium_c"] + context["medium_a"] + context["medium_o"] + context["medium_e"]
        context["all_hard"] = context["hard_c"] + context["hard_a"] + context["hard_o"] + context["hard_e"]

        context["all"] = context["all_simple"] + context["all_medium"] + context["all_hard"]

        if info is None:
            info = {}

        context.update(info)

        template = docxtpl.DocxTemplate(temp_file)
        template.render(context)
        for key, value in context.items():
            if key.lower() == "индикатор" or key.lower() == "компетенция":
                to_file = to_file.replace("[" + key.lower() + "]", str(value).split(" ")[0])
            else:
                to_file = to_file.replace("[" + key.lower() + "]", str(value))
        to_file = to_file.replace("[дата]", f"{datetime.now():%Y_%m_%d %H_%M_%S%z}")
        if not to_file.endswith(".docx") and not to_file.endswith(".doc"):
            to_file += ".docx"
        to_file = to_file.replace("\n", " ")
        template.save(to_file)
        os.remove(temp_file)

        self.path = to_file
        template.save(self.path)

        return self.path


class LayoutStruct:
    questions: dict
    skipped: list
    struct: list[dict] = None
    add_question: dict

    def __init__(self, struct=None):
        self.questions = {}
        self.groups = []
        self.skipped = []
        if struct is not None:
            self.struct = []
            with open(struct, 'r', encoding='utf-8') as f:
                lines = f.readlines()
                headers = lines[0].split(',')
                for line in lines[1:]:
                    values = line.split(',')
                    self.struct.append(
                        {headers[i].strip().lower(): values[i].strip() for i in range(min(len(values), len(headers)))})

                    self.struct[-1]['category'] = int(self.struct[-1]['category'])
                    self.struct[-1]['count'] = int(self.struct[-1]['count'])
                    self.struct[-1]['questions'] = []
                    if self.struct[-1]['group'] not in self.groups:
                        self.groups.append(self.struct[-1]['group'])

    def gift2struct(self, gift_obj: gift.Gift = None, filepath: str = None):
        if gift_obj is not None:
            for q in gift_obj.questions:
                self.add_question(q)
        if filepath is not None:
            gift_obj = gift.from_file(filepath)
            self.gift2struct(gift_obj)

    def add_question(self, question):
        if self.struct is not None:
            category_level = question.get_category_level()
            question_type = question.get_type()[:-2]
            add = False
            for s in self.struct:
                if s['type'] == question_type and category_level == s['category'] and len(s['questions']) < s['count']:
                    add = True
                    s['questions'].append(question)
                    break
            if not add:
                print(f"Question was skipped: {question}")
                self.skipped.append(question)
        else:
            pass

    def get_struct(self):
        return self.struct

    def to_layout(self, layout: Layout) -> Layout:
        layout.set_struct(self.get_struct())
        for struct in self.get_struct():
            for i in range(len(struct['questions'])):
                prefix = struct['prefix'] if i == 0 else None
                postfix = struct['postfix'] if i == len(struct['questions']) - 1 else None
                layout.add(struct['group'], struct['questions'][i], prefix, postfix)
        return layout

    def build(self, template, path):
        layout = Layout(template)
        self.to_layout(layout)
        layout.save(path)
        return layout


def gift2layout(
        questions_file,
        layout_file,
        info=None,
        html_convert=True,
        template_file=None,
        categories_file=None,
        struct_file=None
):
    if info is None:
        info = {}
    if template_file is None:
        template_file = "./template/LayoutTest.docx"
    if categories_file is None:
        categories_file = "./template/categories.csv"
    if struct_file is None:
        struct_file = "./template/struct.csv"

    if isinstance(questions_file, Gift):
        questions = questions_file
    else:
        questions = gift.from_file(filepath=questions_file)

    questions.add_categories(categories=categories_file)
    questions.fix_names()

    # for q in questions.questions:
    #     print(q)

    layout = Layout(template_file, html_convert=html_convert)
    structure = LayoutStruct(struct=struct_file)

    structure.gift2struct(gift_obj=questions)
    structure.to_layout(layout=layout)

    out_filepath = layout.build(to_file=layout_file, info=info)

    return out_filepath, layout, structure, questions
