import re
from copy import deepcopy

import latex2mathml.converter
import markdown
from bs4 import BeautifulSoup
from docx import Document
from docx.document import Document as _Document
from docx.oxml import OxmlElement, CT_Tbl, CT_P
from docx.shared import Inches
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from htmldocx import HtmlToDocx

__STR_FOR_NUMERATION__ = 'num_'
__PATTERN_FOR_NUMERATION__ = "{{ " + __STR_FOR_NUMERATION__ + "{i} }}"
__PATTERN_FOR_NUMERATION_REGEX__ = r"{{\s" + __STR_FOR_NUMERATION__ + r".*\s}}"

from lxml import etree


def contains_html(text):
    return bool(BeautifulSoup(text, "html.parser").find())


def clean_string_regex(s):
    """
    Функция очищает слова в квадратных скобках,
    находящихся в начале строки с помощью регулярных выражений

    :param s: строка для очистки
    :return: очищенная строка
    """
    pattern = r"^\[[a-zA-Z]+\]\s*"
    match = re.search(pattern, s)
    if match:
        # Найдено вхождение паттерна, очищаем строку
        s = s[match.end():].lstrip()
    return s


def iter_block_items(parent):
    """
    Generate a reference to each paragraph and table child within *parent*,
    in document order. Each returned value is an instance of either Table or
    Paragraph. *parent* would most commonly be a reference to a main
    Document object, but also works for a _Cell object, which itself can
    contain paragraphs and tables.
    """
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
        # print(parent_elm.xml)
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("something's not right")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


# def latex_to_word(latex_input):
#     mathml = latex2mathml.converter.convert(latex_input, display="block")
#
#     tree = etree.fromstring(mathml)
#     xslt = etree.parse(
#         './template/MML2OMML.XSL'
#     )
#     transform = etree.XSLT(xslt)
#     new_dom = transform(tree)
#
#     wrapper = etree.fromstring(
#         '''<a14:m xmlns:a14="http://schemas.microsoft.com/office/drawing/2010/main">
#             <m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
#             </m:oMath>
#         </a14:m>''')
#     wrapper.getchildren()[0].append(new_dom.getroot())
#     return wrapper


def text2latex(latex_input):
    mathml = latex2mathml.converter.convert(latex_input)

    tree = etree.fromstring(mathml)
    xslt = etree.parse('./template/MML2OMML.XSL')
    transform = etree.XSLT(xslt)
    new_dom = transform(tree)

    return new_dom.getroot()


def get_latex_equations(text):
    return re.findall(r'\$([^$].*?)\$', text)


def replace_text(text, paragraph, style, text_prepare=None):
    paragraph2 = insert_text_after(text, paragraph, style, text_prepare)
    if paragraph2 != paragraph:
        remove_paragraph(paragraph)
    return paragraph2


def remove_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None


def copy_table_after(table, paragraph):
    tbl, p = table._element, paragraph._element
    new_tbl = deepcopy(tbl)
    p.addnext(new_tbl)
    return Table(new_tbl, p._parent)


def move_table_after(table, paragraph):
    tbl, p = table._element, paragraph._element
    p.addnext(tbl)
    return table


def find_paragraph_id(document, pattern, find, start_id=0):
    test_paragraph_id = -1
    for i in range(start_id, len(document.paragraphs)):
        if find(document.paragraphs[i].text.strip(), pattern):
            test_paragraph_id = i
            break
    return test_paragraph_id


def insert_paragraph_after(paragraph, text=None, style=None):
    new_paragraph = OxmlElement("w:p")
    paragraph._element.addnext(new_paragraph)
    new_para = Paragraph(new_paragraph, paragraph._parent)
    if text:
        new_para.text = text.replace("\\n", "\n").replace("\\t", "\t").replace("\\r", "\r")
    if style:
        new_para.style = style
    return new_para


def insert_html_after(paragraph, html, style=None):
    html = re.sub(r'<img([\w\W]+?)>', "<p>***** ТУТ БЫЛА КАРТИНКА *****</p>", html)
    html = (html
            .replace("\\n", "&nbsp; ")
            .replace("\\t", " ")
            .replace("\\r", "\r"))
    desc_document = Document()
    new_parser = HtmlToDocx()
    new_parser.add_html_to_document(html, desc_document)
    p_and_t = []
    for p in list(iter_block_items(desc_document))[::-1]:
        paragraph._element.addnext(p._element)
        if isinstance(p, Paragraph):
            p_and_t.append(Paragraph(p._element, paragraph._parent))
            # pp = p_and_t[-1]
            # t = pp.text
            # pp.text = ''
            # p_and_t.append(insert_formula_after(pp, t, style, newline=False))
        elif isinstance(p, Table):
            p_and_t.append(Table(p._element, paragraph._parent))
            # for row in p_and_t[-1].rows:
            #     for cell in row.cells:
            #         for pp in cell.paragraphs:
            #             t = pp.text
            #             pp.text = ''
            #             insert_formula_after(pp, t, style, newline=False)

    if style is not None:
        for pt in p_and_t:
            pt.style = style

    return p_and_t


def insert_markdown_after(paragraph, text, style=None):
    paragraphs = insert_html_after(paragraph, markdown.markdown(text))
    if style is not None:
        for paragraph in paragraphs:
            paragraph.style = style
    return paragraphs


def insert_text_after(text, paragraph, style, text_prepare=None):
    is_html = text.strip().startswith("[html]")
    is_markdown = text.strip().startswith("[markdown]")
    is_formula = text.strip().startswith("[formula]")
    n_text = clean_string_regex(text)
    n_text = (n_text
              .replace(r'\:', ':')
              .replace(r'\=', '=')
              .replace(r'\{', '{')
              .replace(r'\}', '}')
              .replace(r'\#', '#'))

    if is_html or is_markdown or is_formula:
        if is_html:
            pt = insert_html_after(paragraph, n_text)
        elif is_formula:
            pt = [insert_formula_after(paragraph, n_text, style)]
        else:
            pt = insert_markdown_after(paragraph, n_text)
        pt = pt[::-1]
        prepared = False
        j = 0
        while j < len(pt):
            is_paragraph = isinstance(pt[j], Paragraph)
            is_table = isinstance(pt[j], Table)
            if is_paragraph:
                if len(pt[j].text.strip()) == 0:
                    remove_paragraph(pt.pop(j))
                    continue
                k = 0
                # print([run.text for run in pt[j].runs])
                for run in reversed(list(pt[j].runs)):
                    if run.text.strip() == "":
                        k += 1
                        pt[j]._p.remove(run._r)
                    else:
                        run.text = run.text.strip(" \n,-")
                        if run.text.endswith("\\:"):
                            run.text = run.text.replace("\\:", ":")
                # print(k)
                if not prepared:
                    if text_prepare is not None:
                        pt[j].text = text_prepare(pt[j].text)
                        prepared = True
                pt[j].style = style
            elif is_table:
                # pt[j].style = style
                for row in pt[j].rows:
                    row.style = style
                    for cell in row.cells:
                        cell.style = style
                        for pp in cell.paragraphs:
                            pp.style = style
                            # replace_text(cell.text, cell.paragraphs[], style, text_prepare)
            j += 1
        if len(pt) > 0:
            paragraph = pt[j - 1]
    elif text.strip().startswith("[moodle]"):
        if text_prepare is not None:
            n_text = text_prepare(n_text)
        paragraph = insert_paragraph_after(paragraph, n_text, style)
    else:
        if text_prepare is not None:
            n_text = text_prepare(n_text)
        paragraph = insert_paragraph_after(paragraph, n_text, style)
    return paragraph


def insert_text_in(paragraph, text, pattern=r'\\\((.*?)\\\)'):
    last_id = 0
    for match in re.finditer(pattern, text):
        start, end = match.span()
        formula = match.group()
        print(f'Найдена формула "{formula}" на позициях с {start} по {end}')
        run = paragraph.add_run(text[last_id:start] + " ")
        try:
            latex_run = text2latex(formula)
            run._element.append(latex_run)
            last_id = end + 1
        except Exception as e:
            print(e)


def add_short_question(paragraph, question, text, styles):
    paragraph = insert_text_after(
        text,
        paragraph,
        styles['question'],
        text_prepare=lambda text: f"{__PATTERN_FOR_NUMERATION__.replace('{i}', question.name.strip())}\t{text}"
    )
    return insert_paragraph_after(paragraph, "", styles['question'])


def add_multiple_question(paragraph, question, text, styles):
    is_html = text.strip().startswith("[html]")
    is_markdown = text.strip().startswith("[markdown]")
    is_moodle = text.strip().startswith("[moodle]")

    paragraph = insert_text_after(
        text,
        paragraph,
        styles['question'],
        text_prepare=lambda text: f"{__PATTERN_FOR_NUMERATION__.replace('{i}', question.name.strip())}\t{text}"
    )
    answers_start = ord("А")
    for answer in question.get_answers():
        prefix = ""
        if answer.text.strip() == clean_string_regex(answer.text.strip()):
            prefix = ""
        elif is_html:
            prefix = "[html]"
        elif is_markdown:
            prefix = "[markdown]"
        elif is_moodle:
            prefix = "[moodle]"
        # print(answer.text)
        paragraph = insert_text_after(
            prefix + answer.text,
            paragraph,
            styles['answer'],
            text_prepare=lambda text: f"{chr(answers_start)})\t{text}"
        )
        answers_start += 1
    return insert_paragraph_after(paragraph, "", styles['question'])


def add_matching_question(document, paragraph, question, text, styles):
    if 'установите соответствие' not in text.lower():
        paragraph = insert_paragraph_after(
            paragraph,
            f"{__PATTERN_FOR_NUMERATION__.replace('{i}', question.name.strip())}\tУстановите соответствие",
            styles['question']
        )
        if not text.startswith("::") or not text.endswith("::"):
            paragraph = insert_text_after(
                text,
                paragraph,
                styles['question'],
            )
    else:
        paragraph = insert_paragraph_after(
            paragraph,
            f"{__PATTERN_FOR_NUMERATION__.replace('{i}', question.name.strip())}\t{text}",
            styles['question']
        )

    table = document.add_table(rows=1, cols=2)
    answers_start = 0
    answers_start_col1 = 1
    answers_start_col2 = ord("А")
    row = None

    is_html = text.strip().startswith("[html]")
    is_markdown = text.strip().startswith("[markdown]")
    is_moodle = text.strip().startswith("[moodle]")

    for answer in question.get_answers():
        if row is None:
            row = table.rows[0]
        else:
            row = table.add_row()
        row.cells[0].width = Inches(3)
        row.cells[1].width = Inches(4.5)
        columns = answer.text.strip().split("->")
        style_ = [styles['table'], styles['key_answer']]
        start_column = [answers_start_col1, answers_start_col2]
        column_func = [lambda x: x, lambda x: chr(x) + ")"]
        prefix = ""
        if answer.text.strip() == clean_string_regex(answer.text.strip()):
            prefix = ""
        elif is_html:
            prefix = "[html]"
        elif is_markdown:
            prefix = "[markdown]"
        elif is_moodle:
            prefix = "[moodle]"

        for k, _ in enumerate(columns):
            if columns[k].strip() == "":
                continue
            replace_text(f"{prefix}{columns[k]}", row.cells[k].paragraphs[0], style_[k])
            r = row.cells[k].paragraphs[0]
            r.text = str(column_func[k](start_column[k] + answers_start)) + " " + r.text.strip()

        answers_start += 1

    paragraph2 = insert_paragraph_after(paragraph, "", styles['question'])
    move_table_after(table, paragraph)
    return paragraph2


def insert_formula_after(paragraph, text, style=None, pattern=r'\\\((.*?)\\\)', newline=True):
    last_id = 0
    for match in re.finditer(pattern, text.replace('\\\\(', '\\(').replace('\\\\)', '\\)')):
        start, end = match.span()
        formula = match.group()
        formula = formula.replace('\\{', '{').replace('\\}', '}')
        formula = formula.replace('\\\\\\', '*TMP_DELIMITER*')
        formula = formula.replace('\\\\', '\\')
        formula = formula.replace('*TMP_DELIMITER*', '\\\\\\')
        print(f'Найдена формула "{formula}" на позициях с {start} по {end}')
        run = paragraph.add_run(text[last_id:start] + " ")
        try:
            latex_element = text2latex(formula[2:-2])
            run._element.addnext(latex_element)
            last_id = end + 1
        except Exception as e:
            print(e)
    run = paragraph.add_run(text[last_id:] + " ")
    if style:
        paragraph.style = style
    if newline:
        return insert_paragraph_after(paragraph, None, style)
    return paragraph
