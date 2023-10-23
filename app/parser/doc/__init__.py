import re
from copy import deepcopy

from .template import fill_test_template
from .editor import clean_string_regex
from docx import Document

import markdown
from docx.document import Document as _Document
from docx.oxml import OxmlElement
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.shared import Inches
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from htmldocx import HtmlToDocx


class DocumentX:

    def __init__(self):
        pass

    @staticmethod
    def iter_block_items(parent):
        """
        Generate a reference to each paragraph and table child within *parent*,
        in document order. Each returned value is an instance of either Table or
        Paragraph. *parent* would most commonly be a reference to a main
        Document object, but also works for a _Cell object, which itself can
        contain paragraphs and tables.
        """
        if isinstance(parent, _Document):
            parent_elm = parent.element.body
            # print(parent_elm.xml)
        elif isinstance(parent, _Cell):
            parent_elm = parent._tc
        else:
            raise ValueError("something's not right")

        for child in parent_elm.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, parent)
            elif isinstance(child, CT_Tbl):
                yield Table(child, parent)

    @staticmethod
    def replace_text(text, paragraph, style, text_prepare=None):
        paragraph2 = DocumentX.insert_text_after(text, paragraph, style, text_prepare)
        if paragraph2 != paragraph:
            DocumentX.remove_paragraph(paragraph)
        return paragraph2

    @staticmethod
    def remove_paragraph(paragraph):
        p = paragraph._element
        p.getparent().remove(p)
        p._p = p._element = None

    @staticmethod
    def copy_table_after(table, paragraph):
        tbl, p = table._element, paragraph._element
        new_tbl = deepcopy(tbl)
        p.addnext(new_tbl)
        return Table(new_tbl, p._parent)

    @staticmethod
    def move_table_after(table, paragraph):
        tbl, p = table._element, paragraph._element
        p.addnext(tbl)
        return table

    @staticmethod
    def find_paragraph_id(document, pattern, find, start_id=0):
        test_paragraph_id = -1
        for i in range(start_id, len(document.paragraphs)):
            if find(document.paragraphs[i].text.strip(), pattern):
                test_paragraph_id = i
                break
        return test_paragraph_id

    @staticmethod
    def insert_paragraph_after(paragraph, text=None, style=None):
        new_paragraph = OxmlElement("w:p")
        paragraph._element.addnext(new_paragraph)
        new_para = Paragraph(new_paragraph, paragraph._parent)
        if text:
            new_para.text = text.replace("\\n", "\n").replace("\\t", "\t").replace("\\r", "\r")
        if style:
            new_para.style = style
        return new_para

    @staticmethod
    def insert_html_after(html, paragraph):
        html = re.sub(r'<img([\w\W]+?)>', "<p>***** ТУТ БЫЛА КАРТИНКА *****</p>", html)
        html = html.replace("\\n", "&nbsp; ").replace("\\t", " ").replace("\\r", "\r")
        desc_document = DocumentX()
        new_parser = HtmlToDocx()
        new_parser.add_html_to_document(html, desc_document)
        p_and_t = []
        for p in list(DocumentX.iter_block_items(desc_document))[::-1]:
            paragraph._element.addnext(p._element)
            if isinstance(p, Paragraph):
                p_and_t.append(Paragraph(p._element, paragraph._parent))
            elif isinstance(p, Table):
                p_and_t.append(Table(p._element, paragraph._parent))
        return p_and_t

    @staticmethod
    def insert_markdown_after(text, paragraph):
        return DocumentX.insert_html_after(markdown.markdown(text), paragraph)

    @staticmethod
    def insert_text_after(text, paragraph, style, text_prepare=None):
        is_html = text.strip().startswith("[html]")
        is_markdown = text.strip().startswith("[markdown]")
        n_text = clean_string_regex(text)
        if is_html or is_markdown:
            if is_html:
                pt = DocumentX.insert_html_after(n_text, paragraph)
            else:
                pt = DocumentX.insert_markdown_after(n_text, paragraph)
            pt = pt[::-1]
            prepared = False
            j = 0
            while j < len(pt):
                is_paragraph = isinstance(pt[j], Paragraph)
                is_table = isinstance(pt[j], Table)
                if is_paragraph:
                    if len(pt[j].text.strip()) == 0:
                        DocumentX.remove_paragraph(pt.pop(j))
                        continue
                    k = 0
                    # print([run.text for run in pt[j].runs])
                    for run in reversed(list(pt[j].runs)):
                        if run.text.strip() == "":
                            k += 1
                            pt[j]._p.remove(run._r)
                        else:
                            run.text = run.text.strip(" \n,-")
                            if run.text.endswith("\\:"):
                                run.text = run.text.replace("\\:", ":")
                    # print(k)
                    if not prepared:
                        if text_prepare is not None:
                            pt[j].text = text_prepare(pt[j].text)
                            prepared = True
                    pt[j].style = style
                elif is_table:
                    # pt[j].style = style
                    for row in pt[j].rows:
                        row.style = style
                        for cell in row.cells:
                            cell.style = style
                            for pp in cell.paragraphs:
                                pp.style = style
                                # replace_text(cell.text, cell.paragraphs[], style, text_prepare)
                j += 1
            if len(pt) > 0:
                paragraph = pt[j - 1]
        elif text.strip().startswith("[moodle]"):
            if text_prepare is not None:
                n_text = text_prepare(n_text)
            paragraph = DocumentX.insert_paragraph_after(paragraph, n_text, style)
        else:
            if text_prepare is not None:
                n_text = text_prepare(n_text)
            paragraph = DocumentX.insert_paragraph_after(paragraph, n_text, style)
        return paragraph
