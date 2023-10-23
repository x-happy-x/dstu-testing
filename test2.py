import os

import docx


def changeAll(old_path, new_path):
    for root, dirs, files in os.walk(old_path):
        for file in files:
            if not file.endswith(".docx"):
                continue
            print(os.path.join(root.replace(old_path, ""), file))
            file_path = os.path.join(root, file)
            doc = docx.Document(file_path)
            if not change(doc):
                continue
            new_file = file_path.replace(old_path, new_path)
            dir_path = root.replace(old_path, new_path)
            if not os.path.exists(dir_path):
                os.makedirs(dir_path)
            doc.save(new_file)


def change(doc):
    t = False
    for table in doc.tables:
        if len(table.rows) < 3 or not table.rows[-1].cells[0].text.strip().lower().startswith("итого") or not \
                table.rows[0].cells[0].text.strip().lower().startswith("компетенция") or table.rows[-2].cells[-2].paragraphs[0].text.strip() == "4":
            continue
        table.rows[-2].cells[-2].paragraphs[0].text = "4"
        t = True
    if not t:
        return False
    return True


changeAll("source/Тесты преподавателей", "source/Тесты преподавателей2")
